"""
Document Parsing Tool — Extract text from PDFs, DOCX, and text files.
"""
import os
import logging
from gateway.tools import Tool, register_tool

logger = logging.getLogger("gateway.tools.document_parse")

MAX_OUTPUT = 15000


class DocumentParseTool(Tool):
    name = "parse_document"
    description = (
        "Extract text content from uploaded documents: PDF, DOCX (Word), "
        "TXT, CSV, JSON, Markdown, and other text files. "
        "Use this to read and analyze document contents."
    )
    parameters = {
        "type": "object",
        "properties": {
            "file_path": {
                "type": "string",
                "description": "Path to the uploaded document file (e.g., /tmp/gateway_uploads/report.pdf).",
            },
            "pages": {
                "type": "string",
                "description": "For PDFs: page range to extract, e.g. '1-5' or '1,3,7'. Default: all pages.",
            },
        },
        "required": ["file_path"],
    }

    async def execute(self, params: dict) -> str:
        file_path = params.get("file_path", "")
        pages_str = params.get("pages", "")

        if not file_path:
            return "Error: file_path is required"
        if not os.path.exists(file_path):
            return f"Error: File not found: {file_path}"

        ext = os.path.splitext(file_path)[1].lower()
        basename = os.path.basename(file_path)

        try:
            if ext == ".pdf":
                text = self._parse_pdf(file_path, pages_str)
            elif ext == ".docx":
                text = self._parse_docx(file_path)
            elif ext in (".txt", ".md", ".csv", ".json", ".xml", ".yaml", ".yml", ".log", ".py", ".js", ".ts", ".html", ".css"):
                with open(file_path, "r", errors="replace") as f:
                    text = f.read()
            else:
                # Try reading as plain text
                try:
                    with open(file_path, "r", errors="replace") as f:
                        text = f.read()
                except Exception:
                    return f"Error: Unsupported file format '{ext}'. Supported: PDF, DOCX, TXT, CSV, JSON, MD, and other text files."

            if len(text) > MAX_OUTPUT:
                text = text[:MAX_OUTPUT] + f"\n\n... (truncated, total {len(text)} chars)"

            logger.info(f"Document parsed: {basename} -> {len(text)} chars")
            return f"## Document: {basename}\n\n{text}"

        except Exception as e:
            logger.exception(f"Document parse error: {file_path}")
            return f"Error parsing document: {str(e)}"

    def _parse_pdf(self, path: str, pages_str: str) -> str:
        page_set = self._parse_page_range(pages_str) if pages_str else None

        # Primary: PyMuPDF (handles most PDFs including non-standard ones)
        try:
            import fitz  # PyMuPDF

            text_parts = []
            doc = fitz.open(path)
            total_pages = len(doc)
            for i in range(total_pages):
                if page_set and (i + 1) not in page_set:
                    continue
                page_text = doc[i].get_text()
                if page_text.strip():
                    text_parts.append(f"--- Page {i + 1}/{total_pages} ---\n{page_text}")
            doc.close()

            if text_parts:
                return "\n\n".join(text_parts)
        except Exception as e:
            logger.warning(f"PyMuPDF failed, trying pdfplumber: {e}")

        # Fallback: pdfplumber
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(path) as pdf:
                total_pages = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    if page_set and (i + 1) not in page_set:
                        continue
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(f"--- Page {i + 1}/{total_pages} ---\n{page_text}")

            if text_parts:
                return "\n\n".join(text_parts)
        except Exception as e:
            logger.warning(f"pdfplumber also failed: {e}")

        return "(No extractable text found in PDF. The document may contain only images or be encrypted.)"

    def _parse_docx(self, path: str) -> str:
        from docx import Document

        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Also extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))

        return "\n\n".join(parts) if parts else "(No text content found in document.)"

    def _parse_page_range(self, s: str) -> set:
        pages = set()
        for part in s.split(","):
            part = part.strip()
            if "-" in part:
                start, end = part.split("-", 1)
                for p in range(int(start), int(end) + 1):
                    pages.add(p)
            else:
                pages.add(int(part))
        return pages


register_tool(DocumentParseTool())
